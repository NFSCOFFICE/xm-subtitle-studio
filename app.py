from __future__ import annotations

import json
import mimetypes
import os
import re
import shutil
import subprocess
import sys
import tempfile
import uuid
import zipfile
from concurrent.futures import ThreadPoolExecutor
from dataclasses import asdict, dataclass, field, fields
from datetime import datetime
from pathlib import Path
from threading import Lock
from typing import Dict, List, Optional

import numpy as np
import librosa
from fastapi import FastAPI, File, Form, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
from fastapi.staticfiles import StaticFiles
from faster_whisper import WhisperModel
from huggingface_hub import snapshot_download
from docx import Document
from sklearn.cluster import AgglomerativeClustering
from transformers import pipeline

try:
    from resemblyzer import VoiceEncoder, preprocess_wav
except Exception:
    VoiceEncoder = None
    preprocess_wav = None


SOURCE_DIR = Path(__file__).resolve().parent
if getattr(sys, "frozen", False):
    APP_DIR = Path(sys.executable).resolve().parent
    RESOURCE_DIR = Path(getattr(sys, "_MEIPASS", APP_DIR))
else:
    APP_DIR = SOURCE_DIR
    RESOURCE_DIR = SOURCE_DIR


def user_data_dir() -> Path:
    if sys.platform == "darwin":
        return Path.home() / "Library" / "Application Support" / "XM Subtitle Studio"
    if sys.platform.startswith("win"):
        root = os.environ.get("LOCALAPPDATA") or os.environ.get("APPDATA")
        if root:
            return Path(root) / "XM Subtitle Studio"
    return Path.home() / ".xm-subtitle-studio"


WRITABLE_DIR = user_data_dir() if getattr(sys, "frozen", False) else APP_DIR
STATIC_DIR = RESOURCE_DIR / "static"
UPLOAD_DIR = WRITABLE_DIR / "uploads"
OUTPUT_DIR = WRITABLE_DIR / "outputs"
MODEL_DIR = RESOURCE_DIR / "models" if (RESOURCE_DIR / "models").exists() else WRITABLE_DIR / "models"
DATA_DIR = WRITABLE_DIR / "data"
JOB_STORE_PATH = DATA_DIR / "jobs.json"
VENDOR_FFMPEG_BIN_CANDIDATES = (
    APP_DIR / "vendor" / "ffmpeg" / "bin",
    APP_DIR.parent / "Frameworks" / "vendor" / "ffmpeg" / "bin",
)
ALLOWED_SUFFIXES = {
    ".mp3",
    ".wav",
    ".m4a",
    ".aac",
    ".flac",
    ".ogg",
    ".mp4",
    ".mov",
    ".mkv",
    ".webm",
    ".m4v",
    ".avi",
}
VIDEO_SUFFIXES = {".mp4", ".mov", ".mkv", ".webm", ".m4v", ".avi"}
SUPPORTED_MEDIA_LABEL = ", ".join(sorted(suffix.lstrip(".") for suffix in ALLOWED_SUFFIXES))

for directory in (STATIC_DIR, UPLOAD_DIR, OUTPUT_DIR, MODEL_DIR, DATA_DIR):
    directory.mkdir(parents=True, exist_ok=True)


def prepend_path_if_exists(path: Path) -> None:
    if not path.exists():
        return
    path_str = str(path)
    path_parts = os.environ.get("PATH", "").split(os.pathsep)
    if path_str not in path_parts:
        os.environ["PATH"] = path_str + os.pathsep + os.environ.get("PATH", "")


for vendor_bin in VENDOR_FFMPEG_BIN_CANDIDATES:
    prepend_path_if_exists(vendor_bin)
for fallback_bin in (
    Path("/opt/homebrew/bin"),
    Path("/usr/local/bin"),
    Path("/opt/local/bin"),
):
    prepend_path_if_exists(fallback_bin)


def _ffmpeg_install_hint() -> str:
    if sys.platform.startswith("win"):
        return (
            "Install options:\n"
            "  - Run start-win.bat once to auto-download FFmpeg into vendor/ffmpeg/\n"
            "  - Or install system-wide: winget install Gyan.FFmpeg"
        )
    if sys.platform == "darwin":
        return "Install via Homebrew: brew install ffmpeg"
    return "Install ffmpeg via your system package manager (e.g., apt install ffmpeg)."


def resolve_ffmpeg_binary(name: str) -> str:
    """Return absolute path to ffmpeg/ffprobe, searching vendor dirs then system PATH."""
    exe_name = f"{name}.exe" if sys.platform.startswith("win") else name
    for vendor_bin in VENDOR_FFMPEG_BIN_CANDIDATES:
        candidate = vendor_bin / exe_name
        if candidate.exists():
            return str(candidate)
    found = shutil.which(name)
    if found:
        return found
    checked = ", ".join(str(d / exe_name) for d in VENDOR_FFMPEG_BIN_CANDIDATES)
    raise RuntimeError(
        f"{name} is required but was not found.\n"
        f"Checked: {checked} and system PATH.\n"
        f"{_ffmpeg_install_hint()}"
    )


@dataclass
class JobState:
    job_id: str
    filename: str
    original_name: str
    language: str
    model_size: str
    translate_to: str
    diarization: bool
    speaker_count: str
    smart_split: bool
    ass_style: str
    status: str = "queued"
    progress: float = 0.0
    message: str = "Waiting to start"
    created_at: str = field(default_factory=lambda: datetime.utcnow().isoformat() + "Z")
    detected_language: Optional[str] = None
    srt_path: Optional[str] = None
    outputs: Dict[str, str] = field(default_factory=dict)
    segments: List[dict] = field(default_factory=list)
    draft_segments: List[dict] = field(default_factory=list)
    draft_updated_at: Optional[str] = None
    error: Optional[str] = None


app = FastAPI(title="Offline Subtitle Studio")
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)
app.mount("/static", StaticFiles(directory=STATIC_DIR), name="static")

jobs: Dict[str, JobState] = {}
jobs_lock = Lock()
executor = ThreadPoolExecutor(max_workers=1)
_model_cache: Dict[str, WhisperModel] = {}
_model_lock = Lock()
_translator_cache: Dict[str, object] = {}
_translator_lock = Lock()
_speaker_encoder: Optional[object] = None
_speaker_encoder_lock = Lock()

TRANSLATION_MODELS = {
    ("zh", "en"): "Helsinki-NLP/opus-mt-zh-en",
    ("en", "zh"): "Helsinki-NLP/opus-mt-en-zh",
    ("ja", "en"): "Helsinki-NLP/opus-mt-ja-en",
    ("ko", "en"): "Helsinki-NLP/opus-mt-ko-en",
    ("fr", "en"): "Helsinki-NLP/opus-mt-fr-en",
    ("de", "en"): "Helsinki-NLP/opus-mt-de-en",
    ("es", "en"): "Helsinki-NLP/opus-mt-es-en",
    ("ru", "en"): "Helsinki-NLP/opus-mt-ru-en",
    ("it", "en"): "Helsinki-NLP/opus-mt-it-en",
}
PUNCTUATION_SPLIT_RE = re.compile(r"(?<=[。！？!?；;：:,，])\s*")
TRANSCRIPTION_LEAD_SILENCE_SECONDS = 0.35
SUBTITLE_START_DELAY_SECONDS = 0.03
STRONG_ALIGNMENT_WINDOW_SECONDS = 1.2
STRONG_ALIGNMENT_MIN_SEGMENT_SECONDS = 0.16
BOUNDARY_ALIGNMENT_MAX_WINDOW_SECONDS = 8.0
BOUNDARY_START_HARD_LIMIT_SECONDS = 0.45
BOUNDARY_END_PAD_SECONDS = 0.05
SPEECH_LOCK_LOOKAROUND_SECONDS = 0.75
SPEECH_LOCK_START_PAD_SECONDS = 0.015
SPEECH_LOCK_END_PAD_SECONDS = 0.08
SPEECH_LOCK_MIN_ACTIVE_SECONDS = 0.12
SPEECH_LOCK_MERGE_GAP_SECONDS = 0.28
SPEECH_LOCK_GROUP_GAP_SECONDS = 1.15
SPEECH_LOCK_FRAME_SECONDS = 0.025
SPEECH_LOCK_HOP_SECONDS = 0.01
SHORT_SEGMENT_WORD_LIMIT = 3
SHORT_SEGMENT_MAX_DURATION_SECONDS = 2.2
SHORT_SEGMENT_BASE_BUFFER_SECONDS = 0.9
SHORT_SEGMENT_CHAR_RATE = 6.5
AUTO_LANGUAGE_PROBE_SECONDS = 45
AUTO_LANGUAGE_MIN_CONFIDENCE = 0.65
WHISPER_VAD_PARAMETERS = {
    "threshold": 0.4,
    "min_speech_duration_ms": 120,
    "min_silence_duration_ms": 450,
    "speech_pad_ms": 120,
}
SPLIT_SEGMENT_MIN_SECONDS = 0.2
SPLIT_SEGMENT_GAP_SECONDS = 0.06
TIMELINE_MIN_SEGMENT_SECONDS = 0.2
TIMELINE_MIN_GAP_SECONDS = 0.08


def serialize_job(job: JobState) -> Dict[str, object]:
    data = asdict(job)
    media_path = UPLOAD_DIR / job.filename
    data["has_draft"] = bool(job.draft_segments)
    data["is_video"] = media_path.suffix.lower() in VIDEO_SUFFIXES
    if media_path.exists():
        data["media_preview_url"] = f"/api/jobs/{job.job_id}/media"
    if job.outputs:
        data["download_urls"] = {
            format_name: f"/api/jobs/{job.job_id}/download/{format_name}"
            for format_name in job.outputs
        }
        data["save_urls"] = {
            format_name: f"/api/jobs/{job.job_id}/save/{format_name}"
            for format_name in job.outputs
        }
        data["save_all_url"] = f"/api/jobs/{job.job_id}/save-all"
        if "srt" in job.outputs:
            data["download_url"] = data["download_urls"]["srt"]
    return data


def persist_jobs(snapshot: Optional[List[Dict[str, object]]] = None) -> None:
    if snapshot is None:
        with jobs_lock:
            snapshot = [asdict(job) for job in jobs.values()]
    JOB_STORE_PATH.write_text(
        json.dumps(snapshot, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )


def load_jobs_from_disk() -> Dict[str, JobState]:
    if not JOB_STORE_PATH.exists():
        return {}

    try:
        payload = json.loads(JOB_STORE_PATH.read_text(encoding="utf-8"))
    except json.JSONDecodeError:
        return {}

    job_fields = {item.name for item in fields(JobState)}
    restored: Dict[str, JobState] = {}
    for record in payload if isinstance(payload, list) else []:
        if not isinstance(record, dict):
            continue
        job_data = {key: value for key, value in record.items() if key in job_fields}
        if not job_data.get("job_id"):
            continue
        job = JobState(**job_data)
        if job.status in {"queued", "running"}:
            job.status = "failed"
            job.message = "Service restarted. Unfinished jobs were converted to drafts. Please submit again."
            job.error = "Interrupted by server restart."
        restored[job.job_id] = job
    return restored


jobs.update(load_jobs_from_disk())
if jobs:
    persist_jobs([asdict(job) for job in jobs.values()])


def update_job(job_id: str, **updates: object) -> None:
    with jobs_lock:
        job = jobs.get(job_id)
        if not job:
            raise HTTPException(status_code=404, detail="Job not found.")
        for key, value in updates.items():
            setattr(job, key, value)
        snapshot = [asdict(item) for item in jobs.values()]
    persist_jobs(snapshot)


def utc_now() -> str:
    return datetime.utcnow().isoformat() + "Z"


def get_audio_duration(file_path: Path) -> float:
    result = subprocess.run(
        [
            resolve_ffmpeg_binary("ffprobe"),
            "-v",
            "error",
            "-show_entries",
            "format=duration",
            "-of",
            "json",
            str(file_path),
        ],
        check=True,
        capture_output=True,
        text=True,
    )
    payload = json.loads(result.stdout or "{}")
    duration = payload.get("format", {}).get("duration", 0)
    return float(duration or 0)


def prepare_audio_for_transcription(file_path: Path) -> tuple[Path, float, Optional[Path]]:
    """Add lead-in silence so VAD/Whisper do not clip first words."""
    if TRANSCRIPTION_LEAD_SILENCE_SECONDS <= 0:
        return file_path, 0.0, None

    temp_dir = Path(tempfile.mkdtemp(prefix="subtitle-transcribe-"))
    padded_path = temp_dir / "lead-padded.wav"
    subprocess.run(
        [
            resolve_ffmpeg_binary("ffmpeg"),
            "-y",
            "-f",
            "lavfi",
            "-t",
            str(TRANSCRIPTION_LEAD_SILENCE_SECONDS),
            "-i",
            "anullsrc=channel_layout=stereo:sample_rate=16000",
            "-i",
            str(file_path),
            "-filter_complex",
            "[1:a]aformat=channel_layouts=stereo,aresample=16000[a1];[0:a][a1]concat=n=2:v=0:a=1[out]",
            "-map",
            "[out]",
            "-ac",
            "2",
            "-ar",
            "16000",
            str(padded_path),
        ],
        check=True,
        capture_output=True,
        text=True,
    )
    return padded_path, TRANSCRIPTION_LEAD_SILENCE_SECONDS, temp_dir


def format_timestamp(seconds: float) -> str:
    milliseconds = max(0, int(round(seconds * 1000)))
    hours, remainder = divmod(milliseconds, 3_600_000)
    minutes, remainder = divmod(remainder, 60_000)
    secs, millis = divmod(remainder, 1_000)
    return f"{hours:02}:{minutes:02}:{secs:02},{millis:03}"


def format_vtt_timestamp(seconds: float) -> str:
    return format_timestamp(seconds).replace(",", ".")


def sanitize_segments(segments: List[dict]) -> List[dict]:
    cleaned: List[dict] = []
    for segment in segments:
        start = max(0.0, float(segment["start"]))
        end = max(start, float(segment["end"]))
        text = str(segment["text"]).strip()
        if not text:
            continue
        item = {"start": start, "end": end, "text": text}
        if segment.get("speaker"):
            item["speaker"] = str(segment["speaker"]).strip()
        if segment.get("translation"):
            item["translation"] = str(segment["translation"]).strip()
        cleaned.append(item)
    return cleaned


def normalize_timeline(segments: List[dict]) -> List[dict]:
    cleaned = sanitize_segments(segments)
    if not cleaned:
        return []

    normalized = [dict(item) for item in sorted(cleaned, key=lambda item: (item["start"], item["end"]))]
    for item in normalized:
        start = max(0.0, float(item["start"]))
        end = max(start + TIMELINE_MIN_SEGMENT_SECONDS, float(item["end"]))
        item["start"] = start
        item["end"] = end

    for index in range(1, len(normalized)):
        previous = normalized[index - 1]
        current = normalized[index]
        required_previous_end = float(current["start"]) - TIMELINE_MIN_GAP_SECONDS
        previous_min_end = float(previous["start"]) + TIMELINE_MIN_SEGMENT_SECONDS
        if float(previous["end"]) > required_previous_end:
            previous["end"] = max(previous_min_end, required_previous_end)

        required_current_start = float(previous["end"]) + TIMELINE_MIN_GAP_SECONDS
        if float(current["start"]) < required_current_start:
            current["start"] = required_current_start
            current["end"] = max(
                float(current["end"]),
                float(current["start"]) + TIMELINE_MIN_SEGMENT_SECONDS,
            )

    for item in normalized:
        item["start"] = round(float(item["start"]), 3)
        item["end"] = round(max(float(item["start"]) + 0.001, float(item["end"])), 3)

    return normalized


def shift_segments(segments: List[dict], offset_seconds: float) -> List[dict]:
    if not offset_seconds:
        return segments
    shifted: List[dict] = []
    for segment in segments:
        item = dict(segment)
        item["start"] = max(0.0, float(item.get("start", 0)) - offset_seconds)
        item["end"] = max(item["start"], float(item.get("end", 0)) - offset_seconds)
        shifted.append(item)
    return shifted


def delay_segment_starts(segments: List[dict], delay_seconds: float) -> List[dict]:
    if delay_seconds <= 0:
        return segments

    delayed: List[dict] = []
    for segment in segments:
        item = dict(segment)
        start = float(item.get("start", 0))
        end = float(item.get("end", 0))
        duration = max(0.0, end - start)
        applied_delay = min(delay_seconds, duration * 0.35)
        item["start"] = max(0.0, round(start + applied_delay, 3))
        item["end"] = max(item["start"], end)
        delayed.append(item)
    return delayed


def align_segment_boundaries_to_audio(segments: List[dict], audio_path: Path) -> List[dict]:
    if not segments:
        return segments

    try:
        waveform, sample_rate = librosa.load(str(audio_path), sr=16000, mono=True)
    except Exception:
        return segments

    if waveform.size == 0:
        return segments

    frame_length = max(1, int(sample_rate * 0.02))
    hop_length = max(1, int(sample_rate * 0.01))
    aligned_segments: List[dict] = []

    for segment in segments:
        item = dict(segment)
        start = float(item.get("start", 0.0))
        end = float(item.get("end", 0.0))
        duration = max(0.0, end - start)
        if duration < STRONG_ALIGNMENT_MIN_SEGMENT_SECONDS:
            aligned_segments.append(item)
            continue

        search_start = max(0.0, start)
        search_end = min(end, start + BOUNDARY_ALIGNMENT_MAX_WINDOW_SECONDS)
        if search_end - search_start < 0.04:
            aligned_segments.append(item)
            continue

        start_sample = int(search_start * sample_rate)
        end_sample = int(search_end * sample_rate)
        audio_slice = waveform[start_sample:end_sample]
        if audio_slice.size < frame_length:
            aligned_segments.append(item)
            continue

        rms = librosa.feature.rms(
            y=audio_slice,
            frame_length=frame_length,
            hop_length=hop_length,
            center=False,
        )[0]
        if rms.size == 0:
            aligned_segments.append(item)
            continue

        smoothed_rms = np.convolve(rms, np.ones(3, dtype=np.float32) / 3, mode="same")
        peak = float(np.percentile(smoothed_rms, 95))
        floor = float(np.percentile(smoothed_rms, 25))
        if peak <= 1e-4 or peak <= floor:
            aligned_segments.append(item)
            continue

        threshold = max(floor * 2.6, floor + (peak - floor) * 0.48, 0.006)
        active_frames = np.flatnonzero(smoothed_rms >= threshold)
        if active_frames.size == 0:
            aligned_segments.append(item)
            continue

        first_active = int(active_frames[0])
        last_active = int(active_frames[-1])

        detected_start = search_start + (first_active * hop_length) / sample_rate
        detected_end = search_start + ((last_active * hop_length) + frame_length) / sample_rate
        bounded_start = min(max(start, detected_start), start + BOUNDARY_START_HARD_LIMIT_SECONDS)
        bounded_end = max(bounded_start + 0.04, min(end, detected_end + BOUNDARY_END_PAD_SECONDS))

        if bounded_start > start:
            item["start"] = round(bounded_start, 3)
        if bounded_end < end:
            item["end"] = round(bounded_end, 3)

        aligned_segments.append(item)

    return aligned_segments


def detect_speech_ranges(audio_path: Path) -> List[tuple[float, float]]:
    """Detect coarse speech/audio-active islands used to prevent subtitles over silence."""
    try:
        waveform, sample_rate = librosa.load(str(audio_path), sr=16000, mono=True)
    except Exception:
        return []

    if waveform.size == 0:
        return []

    frame_length = max(1, int(sample_rate * SPEECH_LOCK_FRAME_SECONDS))
    hop_length = max(1, int(sample_rate * SPEECH_LOCK_HOP_SECONDS))
    rms = librosa.feature.rms(
        y=waveform,
        frame_length=frame_length,
        hop_length=hop_length,
        center=False,
    )[0]
    if rms.size == 0:
        return []

    smoothed = np.convolve(rms, np.ones(5, dtype=np.float32) / 5, mode="same")
    floor = float(np.percentile(smoothed, 20))
    body = float(np.percentile(smoothed, 75))
    peak = float(np.percentile(smoothed, 96))
    if peak <= 1e-5:
        return []

    threshold = max(0.0045, floor * 2.4, floor + (body - floor) * 0.7, peak * 0.10)
    active_frames = np.flatnonzero(smoothed >= threshold)
    if active_frames.size == 0:
        return []

    ranges: List[tuple[float, float]] = []
    range_start = int(active_frames[0])
    previous_frame = int(active_frames[0])
    max_gap_frames = max(1, int(SPEECH_LOCK_MERGE_GAP_SECONDS / SPEECH_LOCK_HOP_SECONDS))

    for frame in active_frames[1:]:
        frame_index = int(frame)
        if frame_index - previous_frame > max_gap_frames:
            start_time = range_start * hop_length / sample_rate
            end_time = (previous_frame * hop_length + frame_length) / sample_rate
            if end_time - start_time >= SPEECH_LOCK_MIN_ACTIVE_SECONDS:
                ranges.append((start_time, end_time))
            range_start = frame_index
        previous_frame = frame_index

    start_time = range_start * hop_length / sample_rate
    end_time = (previous_frame * hop_length + frame_length) / sample_rate
    if end_time - start_time >= SPEECH_LOCK_MIN_ACTIVE_SECONDS:
        ranges.append((start_time, end_time))

    return ranges


def lock_segments_to_speech_ranges(segments: List[dict], audio_path: Path) -> List[dict]:
    speech_ranges = detect_speech_ranges(audio_path)
    if not segments or not speech_ranges:
        return segments

    locked_segments: List[dict] = []
    for segment in segments:
        item = dict(segment)
        start = float(item.get("start", 0.0))
        end = float(item.get("end", start))
        if end <= start:
            locked_segments.append(item)
            continue

        search_start = max(0.0, start - SPEECH_LOCK_LOOKAROUND_SECONDS)
        search_end = end + SPEECH_LOCK_LOOKAROUND_SECONDS
        matches = [
            (speech_start, speech_end)
            for speech_start, speech_end in speech_ranges
            if speech_end >= search_start and speech_start <= search_end
        ]
        if not matches:
            locked_segments.append(item)
            continue

        groups: List[List[tuple[float, float]]] = []
        for speech_range in matches:
            if groups and speech_range[0] - groups[-1][-1][1] <= SPEECH_LOCK_GROUP_GAP_SECONDS:
                groups[-1].append(speech_range)
            else:
                groups.append([speech_range])

        def group_score(group: List[tuple[float, float]]) -> float:
            active_duration = sum(
                max(0.0, min(end, speech_end) - max(start, speech_start))
                for speech_start, speech_end in group
            )
            group_start = group[0][0]
            group_end = group[-1][1]
            distance_penalty = min(abs(group_start - start), abs(group_end - end)) * 0.08
            return active_duration - distance_penalty

        best_group = max(groups, key=group_score)
        active_start = best_group[0][0]
        active_end = best_group[-1][1]

        next_start = max(start, active_start - SPEECH_LOCK_START_PAD_SECONDS)
        next_end = min(end, active_end + SPEECH_LOCK_END_PAD_SECONDS)

        if next_end - next_start < TIMELINE_MIN_SEGMENT_SECONDS:
            nearest = max(
                matches,
                key=lambda item_range: min(abs(item_range[0] - start), abs(item_range[1] - end)),
            )
            next_start = max(start, nearest[0] - SPEECH_LOCK_START_PAD_SECONDS)
            next_end = min(end, nearest[1] + SPEECH_LOCK_END_PAD_SECONDS)

        if next_end - next_start >= TIMELINE_MIN_SEGMENT_SECONDS:
            item["start"] = round(next_start, 3)
            item["end"] = round(next_end, 3)

        locked_segments.append(item)

    return locked_segments


def resolve_segment_start(segment: object) -> float:
    base_start = float(getattr(segment, "start", 0.0) or 0.0)
    words = list(getattr(segment, "words", None) or [])
    if not words:
        return base_start

    fallback_start: Optional[float] = None
    for word in words:
        word_start = getattr(word, "start", None)
        if word_start is None:
            continue
        candidate_start = float(word_start)
        if fallback_start is None:
            fallback_start = candidate_start
        token = str(getattr(word, "word", "") or "").strip()
        if any(char.isalnum() for char in token):
            return max(base_start, candidate_start)

    return max(base_start, fallback_start if fallback_start is not None else base_start)


def resolve_segment_end(segment: object, resolved_start: float) -> float:
    base_end = float(getattr(segment, "end", resolved_start) or resolved_start)
    words = list(getattr(segment, "words", None) or [])
    if not words:
        return max(base_end, resolved_start + 0.04)

    fallback_end: Optional[float] = None
    for word in reversed(words):
        word_end = getattr(word, "end", None)
        if word_end is None:
            continue
        candidate_end = float(word_end)
        if fallback_end is None:
            fallback_end = candidate_end
        token = str(getattr(word, "word", "") or "").strip()
        if any(char.isalnum() for char in token):
            return max(resolved_start + 0.04, min(base_end, candidate_end))

    if fallback_end is not None:
        return max(resolved_start + 0.04, min(base_end, fallback_end))
    return max(base_end, resolved_start + 0.04)


def tighten_sparse_long_segments(segments: List[dict]) -> List[dict]:
    if not segments:
        return segments

    tightened: List[dict] = [dict(item) for item in segments]
    for index, segment in enumerate(tightened):
        if index + 1 >= len(tightened):
            continue

        start = float(segment.get("start", 0.0))
        end = float(segment.get("end", 0.0))
        duration = max(0.0, end - start)
        if duration <= SHORT_SEGMENT_MAX_DURATION_SECONDS:
            continue

        text = str(segment.get("text", "")).strip()
        tokens = [token for token in re.split(r"\s+", text) if token]
        if len(tokens) == 0 or len(tokens) > SHORT_SEGMENT_WORD_LIMIT:
            continue

        char_count = len(re.sub(r"\s+", "", text))
        predicted_max = SHORT_SEGMENT_BASE_BUFFER_SECONDS + (char_count / SHORT_SEGMENT_CHAR_RATE)
        predicted_max = max(1.2, min(predicted_max, SHORT_SEGMENT_MAX_DURATION_SECONDS))
        if duration <= predicted_max + 0.15:
            continue

        capped_end = start + predicted_max
        next_start = float(tightened[index + 1].get("start", capped_end))
        capped_end = min(capped_end, next_start - 0.05)

        if capped_end - start >= TIMELINE_MIN_SEGMENT_SECONDS:
            segment["end"] = round(capped_end, 3)

    return tightened


def split_text_for_subtitles(text: str, max_chars: int = 26) -> List[str]:
    normalized = re.sub(r"\s+", " ", text).strip()
    if not normalized:
        return []

    chunks: List[str] = []
    sentence_parts = [part.strip() for part in PUNCTUATION_SPLIT_RE.split(normalized) if part.strip()]
    if not sentence_parts:
        sentence_parts = [normalized]

    for part in sentence_parts:
        if len(part) <= max_chars:
            chunks.append(part)
            continue

        words = part.split(" ")
        if len(words) == 1:
            start = 0
            while start < len(part):
                chunks.append(part[start : start + max_chars].strip())
                start += max_chars
            continue

        current = ""
        for word in words:
            candidate = f"{current} {word}".strip()
            if current and len(candidate) > max_chars:
                chunks.append(current)
                current = word
            else:
                current = candidate
        if current:
            chunks.append(current)

    return [chunk for chunk in chunks if chunk]


def optimize_subtitle_segments(segments: List[dict]) -> List[dict]:
    optimized: List[dict] = []
    for segment in segments:
        parts = split_text_for_subtitles(segment["text"])
        if len(parts) <= 1:
            optimized.append(segment)
            continue

        duration = max(SPLIT_SEGMENT_MIN_SECONDS, segment["end"] - segment["start"])
        total_chars = sum(max(1, len(part)) for part in parts)
        cursor = segment["start"]
        total_gap = SPLIT_SEGMENT_GAP_SECONDS * max(0, len(parts) - 1)
        distributable_duration = max(
            SPLIT_SEGMENT_MIN_SECONDS * len(parts),
            duration - total_gap,
        )

        for index, part in enumerate(parts):
            weight = max(1, len(part)) / total_chars
            piece_duration = distributable_duration * weight
            end = (
                segment["end"]
                if index == len(parts) - 1
                else min(segment["end"], cursor + piece_duration)
            )
            item = {
                "start": round(cursor, 3),
                "end": round(max(cursor + SPLIT_SEGMENT_MIN_SECONDS, end), 3),
                "text": part,
            }
            if segment.get("speaker"):
                item["speaker"] = segment["speaker"]
            optimized.append(item)
            cursor = min(segment["end"], item["end"] + SPLIT_SEGMENT_GAP_SECONDS)

    return normalize_timeline(optimized)


def render_segment_text(segment: dict, bilingual: bool = False) -> str:
    text = segment["text"].strip()
    translation = str(segment.get("translation", "")).strip()
    speaker = str(segment.get("speaker", "")).strip()
    if speaker:
        text = f"[{speaker}] {text}"
    if bilingual and translation:
        return f"{text}\n{translation}"
    return text


def write_srt(segments: List[dict], output_path: Path, bilingual: bool = False) -> None:
    lines: List[str] = []
    for index, segment in enumerate(segments, start=1):
        lines.append(str(index))
        lines.append(
            f"{format_timestamp(segment['start'])} --> {format_timestamp(segment['end'])}"
        )
        lines.append(render_segment_text(segment, bilingual=bilingual))
        lines.append("")
    output_path.write_text("\r\n".join(lines), encoding="utf-8")


def write_vtt(segments: List[dict], output_path: Path, bilingual: bool = False) -> None:
    lines: List[str] = ["WEBVTT", ""]
    for segment in segments:
        lines.append(
            f"{format_vtt_timestamp(segment['start'])} --> {format_vtt_timestamp(segment['end'])}"
        )
        lines.append(render_segment_text(segment, bilingual=bilingual))
        lines.append("")
    output_path.write_text("\r\n".join(lines), encoding="utf-8")


def write_txt(segments: List[dict], output_path: Path, bilingual: bool = False) -> None:
    lines = [
        render_segment_text(segment, bilingual=bilingual)
        for segment in segments
        if segment["text"].strip()
    ]
    output_path.write_text("\n".join(lines), encoding="utf-8")


def write_json_transcript(
    segments: List[dict], output_path: Path, bilingual: bool = False
) -> None:
    output_path.write_text(
        json.dumps({"bilingual": bilingual, "segments": segments}, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )


def write_markdown_transcript(
    segments: List[dict], output_path: Path, bilingual: bool = False
) -> None:
    lines: List[str] = ["# Transcript", ""]
    for segment in segments:
        start = format_timestamp(segment["start"])
        end = format_timestamp(segment["end"])
        speaker = str(segment.get("speaker", "")).strip()
        label = f"**[{start} - {end}]**"
        if speaker:
            label += f" **{speaker}**"
        lines.append(label)
        lines.append("")
        lines.append(segment["text"].strip())
        if bilingual and segment.get("translation"):
            lines.append("")
            lines.append(f"> {str(segment['translation']).strip()}")
        lines.append("")
    output_path.write_text("\n".join(lines), encoding="utf-8")


def write_docx_transcript(
    segments: List[dict], output_path: Path, bilingual: bool = False
) -> None:
    document = Document()
    document.add_heading("Transcript", level=1)
    for segment in segments:
        heading = f"{format_timestamp(segment['start'])} - {format_timestamp(segment['end'])}"
        if segment.get("speaker"):
            heading += f"  {segment['speaker']}"
        document.add_paragraph(heading, style="Heading 3")
        document.add_paragraph(segment["text"].strip())
        if bilingual and segment.get("translation"):
            document.add_paragraph(str(segment["translation"]).strip())
    document.save(output_path)


def format_ass_timestamp(seconds: float) -> str:
    total = max(0, int(round(seconds * 100)))
    hours, remainder = divmod(total, 360000)
    minutes, remainder = divmod(remainder, 6000)
    secs, centis = divmod(remainder, 100)
    return f"{hours}:{minutes:02}:{secs:02}.{centis:02}"


def ass_style_block(_style_name: str) -> str:
    return "Style: Main,Arial,24,&H00FFFFFF,&H000000FF,&H0010182A,&H8010182A,0,0,0,0,100,100,0,0,1,2,0,2,24,24,18,1"


def write_ass_subtitles(
    segments: List[dict], output_path: Path, style_name: str, bilingual: bool = False
) -> None:
    header = "\n".join(
        [
            "[Script Info]",
            "Title: Offline Subtitle Studio",
            "ScriptType: v4.00+",
            "WrapStyle: 2",
            "ScaledBorderAndShadow: yes",
            "",
            "[V4+ Styles]",
            "Format: Name,Fontname,Fontsize,PrimaryColour,SecondaryColour,OutlineColour,BackColour,Bold,Italic,Underline,StrikeOut,ScaleX,ScaleY,Spacing,Angle,BorderStyle,Outline,Shadow,Alignment,MarginL,MarginR,MarginV,Encoding",
            ass_style_block(style_name),
            "",
            "[Events]",
            "Format: Layer,Start,End,Style,Name,MarginL,MarginR,MarginV,Effect,Text",
        ]
    )

    lines = [header]
    for segment in segments:
        text = render_segment_text(segment, bilingual=bilingual).replace("\n", r"\N")
        lines.append(
            f"Dialogue: 0,{format_ass_timestamp(segment['start'])},{format_ass_timestamp(segment['end'])},Main,,0,0,0,,{text}"
        )
    output_path.write_text("\n".join(lines), encoding="utf-8")


_FW_REPO_MAP = {
    "tiny": "Systran/faster-whisper-tiny",
    "base": "Systran/faster-whisper-base",
    "small": "Systran/faster-whisper-small",
    "medium": "Systran/faster-whisper-medium",
    "large-v1": "Systran/faster-whisper-large-v1",
    "large-v2": "Systran/faster-whisper-large-v2",
    "large-v3": "Systran/faster-whisper-large-v3",
}


def _ensure_local_whisper_model(model_size: str) -> str:
    # If caller passed a local path, honor it.
    candidate = Path(model_size)
    if candidate.exists():
        return str(candidate)
    repo_id = _FW_REPO_MAP.get(model_size, f"Systran/faster-whisper-{model_size}")
    local_dir = MODEL_DIR / repo_id.split("/")[-1]
    if not (local_dir / "model.bin").exists():
        local_dir.mkdir(parents=True, exist_ok=True)
        # local_dir avoids the HF cache's blob+symlink layout, which fails on
        # Windows without SeCreateSymbolicLinkPrivilege (WinError 1314).
        snapshot_download(repo_id=repo_id, local_dir=str(local_dir))
    return str(local_dir)


def load_model(model_size: str) -> WhisperModel:
    with _model_lock:
        if model_size not in _model_cache:
            _model_cache[model_size] = WhisperModel(
                _ensure_local_whisper_model(model_size),
                device="cpu",
                compute_type="int8",
            )
        return _model_cache[model_size]


def probe_auto_language(model: WhisperModel, audio_path: Path) -> tuple[Optional[str], Optional[float]]:
    temp_dir = Path(tempfile.mkdtemp(prefix="subtitle-lang-probe-"))
    probe_path = temp_dir / "lang-probe.wav"
    try:
        subprocess.run(
            [
                resolve_ffmpeg_binary("ffmpeg"),
                "-y",
                "-i",
                str(audio_path),
                "-t",
                str(AUTO_LANGUAGE_PROBE_SECONDS),
                "-ac",
                "1",
                "-ar",
                "16000",
                str(probe_path),
            ],
            check=True,
            capture_output=True,
            text=True,
        )
        _segments, info = model.transcribe(
            str(probe_path),
            language=None,
            vad_filter=True,
            vad_parameters=WHISPER_VAD_PARAMETERS,
            beam_size=1,
            language_detection_segments=3,
            word_timestamps=False,
        )
        language = getattr(info, "language", None)
        probability = getattr(info, "language_probability", None)
        if probability is not None:
            try:
                probability = float(probability)
            except (TypeError, ValueError):
                probability = None
        return language, probability
    except Exception:
        return None, None
    finally:
        shutil.rmtree(temp_dir, ignore_errors=True)


def load_speaker_encoder() -> object:
    global _speaker_encoder
    if VoiceEncoder is None:
        raise RuntimeError(
            "Speaker diarization requires optional dependencies. "
            "Install them with: python -m pip install -r requirements-speaker.txt"
        )
    with _speaker_encoder_lock:
        if _speaker_encoder is None:
            _speaker_encoder = VoiceEncoder()
        return _speaker_encoder


def get_translation_pipeline(source_language: str, target_language: str):
    model_name = TRANSLATION_MODELS.get((source_language, target_language))
    if not model_name:
        return None

    cache_key = f"{source_language}:{target_language}"
    with _translator_lock:
        if cache_key not in _translator_cache:
            _translator_cache[cache_key] = pipeline(
                "translation",
                model=model_name,
                tokenizer=model_name,
                device=-1,
            )
        return _translator_cache[cache_key]


def translate_segments(
    segments: List[dict], source_language: str, target_language: str
) -> List[dict]:
    if not target_language or target_language == "none" or source_language == target_language:
        return segments

    translator = get_translation_pipeline(source_language, target_language)
    if translator is None:
        return segments

    texts = [segment["text"] for segment in segments]
    translations = translator(texts, batch_size=8, max_length=512)
    translated_segments: List[dict] = []
    for segment, translated in zip(segments, translations):
        item = dict(segment)
        item["translation"] = translated["translation_text"].strip()
        translated_segments.append(item)
    return translated_segments


def convert_audio_for_diarization(audio_path: Path) -> Path:
    temp_dir = Path(tempfile.mkdtemp(prefix="subtitle-diarize-"))
    wav_path = temp_dir / f"{audio_path.stem}.wav"
    subprocess.run(
        [
            resolve_ffmpeg_binary("ffmpeg"),
            "-y",
            "-i",
            str(audio_path),
            "-ac",
            "1",
            "-ar",
            "16000",
            str(wav_path),
        ],
        check=True,
        capture_output=True,
        text=True,
    )
    return wav_path


def apply_speaker_diarization(
    segments: List[dict], audio_path: Path, speaker_count: str
) -> List[dict]:
    usable_segments = [segment for segment in segments if (segment["end"] - segment["start"]) >= 0.8]
    if len(usable_segments) < 2:
        return segments
    if preprocess_wav is None:
        raise RuntimeError(
            "Speaker diarization requires optional dependencies. "
            "Install them with: python -m pip install -r requirements-speaker.txt"
        )

    wav_path = convert_audio_for_diarization(audio_path)
    try:
        wav = preprocess_wav(str(wav_path))
        sample_rate = 16000
        encoder = load_speaker_encoder()

        embeddings = []
        usable_indexes = []
        for index, segment in enumerate(segments):
            start_frame = int(segment["start"] * sample_rate)
            end_frame = int(segment["end"] * sample_rate)
            if end_frame - start_frame < int(0.8 * sample_rate):
                continue
            clip = wav[start_frame:end_frame]
            if clip.size < int(0.8 * sample_rate):
                continue
            embeddings.append(encoder.embed_utterance(clip))
            usable_indexes.append(index)

        if len(embeddings) < 2:
            return segments

        n_clusters = int(speaker_count) if speaker_count.isdigit() else min(2, len(embeddings))
        n_clusters = max(2, min(n_clusters, len(embeddings)))
        clustering = AgglomerativeClustering(n_clusters=n_clusters)
        labels = clustering.fit_predict(np.vstack(embeddings))

        speaker_map = {index: f"SPK {label + 1}" for index, label in zip(usable_indexes, labels)}
        diarized_segments = []
        last_speaker = "SPK 1"
        for index, segment in enumerate(segments):
            speaker = speaker_map.get(index, last_speaker)
            item = dict(segment)
            item["speaker"] = speaker
            diarized_segments.append(item)
            last_speaker = speaker
        return diarized_segments
    finally:
        shutil.rmtree(wav_path.parent, ignore_errors=True)


def build_output_filename(original_name: str) -> str:
    original_stem = Path(original_name).stem.strip() or "subtitle"
    safe_stem = re.sub(r'[<>:"/\\|?*\x00-\x1f]+', "-", original_stem)
    safe_stem = re.sub(r"\s+", " ", safe_stem).strip(" .") or "subtitle"
    return f"{safe_stem}.srt"


def reserve_output_paths(original_name: str) -> Dict[str, Path]:
    base_name = build_output_filename(original_name)
    stem = Path(base_name).stem
    candidate = stem
    counter = 1

    while True:
        paths = {
            "srt": OUTPUT_DIR / f"{candidate}.srt",
            "vtt": OUTPUT_DIR / f"{candidate}.vtt",
            "txt": OUTPUT_DIR / f"{candidate}.txt",
            "json": OUTPUT_DIR / f"{candidate}.json",
            "md": OUTPUT_DIR / f"{candidate}.md",
            "docx": OUTPUT_DIR / f"{candidate}.docx",
            "ass": OUTPUT_DIR / f"{candidate}.ass",
            "srt_bilingual": OUTPUT_DIR / f"{candidate}.bilingual.srt",
            "vtt_bilingual": OUTPUT_DIR / f"{candidate}.bilingual.vtt",
            "txt_bilingual": OUTPUT_DIR / f"{candidate}.bilingual.txt",
            "json_bilingual": OUTPUT_DIR / f"{candidate}.bilingual.json",
            "md_bilingual": OUTPUT_DIR / f"{candidate}.bilingual.md",
            "docx_bilingual": OUTPUT_DIR / f"{candidate}.bilingual.docx",
            "ass_bilingual": OUTPUT_DIR / f"{candidate}.bilingual.ass",
        }
        if not any(path.exists() for path in paths.values()):
            return paths
        candidate = f"{stem}-{counter}"
        counter += 1


def write_all_outputs(segments: List[dict], output_paths: Dict[str, Path], ass_style: str) -> List[dict]:
    normalized_segments = normalize_timeline(segments)
    write_srt(normalized_segments, output_paths["srt"])
    write_vtt(normalized_segments, output_paths["vtt"])
    write_txt(normalized_segments, output_paths["txt"])
    write_json_transcript(normalized_segments, output_paths["json"])
    write_markdown_transcript(normalized_segments, output_paths["md"])
    write_docx_transcript(normalized_segments, output_paths["docx"])
    write_ass_subtitles(normalized_segments, output_paths["ass"], ass_style)
    if any(segment.get("translation") for segment in normalized_segments):
        write_srt(normalized_segments, output_paths["srt_bilingual"], bilingual=True)
        write_vtt(normalized_segments, output_paths["vtt_bilingual"], bilingual=True)
        write_txt(normalized_segments, output_paths["txt_bilingual"], bilingual=True)
        write_json_transcript(normalized_segments, output_paths["json_bilingual"], bilingual=True)
        write_markdown_transcript(normalized_segments, output_paths["md_bilingual"], bilingual=True)
        write_docx_transcript(normalized_segments, output_paths["docx_bilingual"], bilingual=True)
        write_ass_subtitles(normalized_segments, output_paths["ass_bilingual"], ass_style, bilingual=True)
    return normalized_segments


def existing_output_map(output_paths: Dict[str, Path]) -> Dict[str, str]:
    return {name: str(path) for name, path in output_paths.items() if path.exists()}


def user_downloads_dir() -> Path:
    downloads = Path.home() / "Downloads"
    downloads.mkdir(parents=True, exist_ok=True)
    return downloads


def unique_destination_path(directory: Path, filename: str) -> Path:
    safe_name = re.sub(r'[<>:"/\\|?*\x00-\x1f]+', "-", filename).strip(" .") or "download"
    candidate = directory / safe_name
    if not candidate.exists():
        return candidate

    stem = candidate.stem
    suffix = candidate.suffix
    counter = 1
    while True:
        numbered = directory / f"{stem}-{counter}{suffix}"
        if not numbered.exists():
            return numbered
        counter += 1


def copy_to_downloads(source_path: Path) -> Path:
    if not source_path.exists():
        raise HTTPException(status_code=404, detail="File not found.")
    destination = unique_destination_path(user_downloads_dir(), source_path.name)
    shutil.copy2(source_path, destination)
    return destination


def build_bundle_zip(job: JobState) -> Path:
    bundle_path = OUTPUT_DIR / f"{Path(build_output_filename(job.original_name)).stem}.bundle.zip"
    counter = 1
    while bundle_path.exists():
        bundle_path = OUTPUT_DIR / f"{Path(build_output_filename(job.original_name)).stem}-{counter}.bundle.zip"
        counter += 1

    with zipfile.ZipFile(bundle_path, "w", compression=zipfile.ZIP_DEFLATED) as archive:
        for path in sorted(Path(file_path) for file_path in job.outputs.values() if Path(file_path).exists()):
            archive.write(path, arcname=path.name)
    return bundle_path


def transcribe_job(
    job_id: str,
    audio_path: Path,
    language: str,
    model_size: str,
    translate_to: str,
    diarization: bool,
    speaker_count: str,
    smart_split: bool,
    ass_style: str,
) -> None:
    transcription_audio_path = audio_path
    lead_offset = 0.0
    transcription_temp_dir: Optional[Path] = None
    try:
        update_job(job_id, status="running", progress=0.05, message="Loading model")
        model = load_model(model_size)

        total_duration = get_audio_duration(audio_path)
        transcription_audio_path, lead_offset, transcription_temp_dir = prepare_audio_for_transcription(audio_path)
        task_language = None if language == "auto" else language
        detected_language: Optional[str] = None
        if language == "auto":
            update_job(job_id, progress=0.11, message="Detecting language")
            probed_language, probed_confidence = probe_auto_language(model, transcription_audio_path)
            if probed_language and (
                probed_confidence is None or probed_confidence >= AUTO_LANGUAGE_MIN_CONFIDENCE
            ):
                task_language = probed_language
                detected_language = probed_language

        update_job(job_id, progress=0.15, message="Running offline transcription")
        segments_iter, info = model.transcribe(
            str(transcription_audio_path),
            language=task_language,
            vad_filter=True,
            vad_parameters=WHISPER_VAD_PARAMETERS,
            beam_size=5,
            language_detection_segments=3,
            word_timestamps=True,
        )

        segments: List[dict] = []
        detected_language = getattr(info, "language", None) or detected_language
        update_job(
            job_id,
            detected_language=detected_language,
            message="Building subtitle timeline",
        )

        for segment in segments_iter:
            resolved_start = resolve_segment_start(segment)
            resolved_end = resolve_segment_end(segment, resolved_start)
            item = {
                "start": resolved_start,
                "end": resolved_end,
                "text": segment.text.strip(),
            }
            if item["text"]:
                segments.append(item)
            if total_duration > 0:
                progress = min(0.95, 0.15 + (item["end"] / total_duration) * 0.75)
                update_job(job_id, progress=progress)

        if not segments:
            raise RuntimeError("No speech was detected in the uploaded audio.")

        segments = shift_segments(segments, lead_offset)
        segments = align_segment_boundaries_to_audio(segments, audio_path)
        segments = lock_segments_to_speech_ranges(segments, audio_path)
        segments = sanitize_segments(delay_segment_starts(segments, SUBTITLE_START_DELAY_SECONDS))
        segments = tighten_sparse_long_segments(segments)
        if smart_split:
            update_job(job_id, progress=0.78, message="Optimizing subtitle breaks")
            segments = optimize_subtitle_segments(segments)
        if diarization:
            update_job(job_id, progress=0.82, message="Detecting speakers")
            segments = apply_speaker_diarization(segments, audio_path, speaker_count)
        segments = normalize_timeline(segments)
        source_language = detected_language or (language if language != "auto" else "")
        update_job(job_id, progress=0.9, message="Preparing exports")
        translated_segments = translate_segments(segments, source_language, translate_to)
        with jobs_lock:
            job = jobs[job_id]
            output_paths = reserve_output_paths(job.original_name)

        update_job(job_id, progress=0.97, message="Writing subtitle files")
        translated_segments = write_all_outputs(translated_segments, output_paths, ass_style)

        update_job(
            job_id,
            status="completed",
            progress=1.0,
            message="Subtitle file is ready",
            srt_path=str(output_paths["srt"]),
            outputs=existing_output_map(output_paths),
            segments=translated_segments,
        )
    except Exception as exc:  # noqa: BLE001
        update_job(
            job_id,
            status="failed",
            message="Transcription failed",
            error=str(exc),
        )
    finally:
        if transcription_temp_dir:
            shutil.rmtree(transcription_temp_dir, ignore_errors=True)


@app.get("/")
def index() -> FileResponse:
    return FileResponse(STATIC_DIR / "index.html")


@app.get("/favicon.ico")
def favicon() -> FileResponse:
    return FileResponse(STATIC_DIR / "favicon.svg", media_type="image/svg+xml")


@app.get("/api/jobs")
def list_jobs() -> List[Dict[str, object]]:
    with jobs_lock:
        return [
            serialize_job(job)
            for job in sorted(
                jobs.values(),
                key=lambda item: item.created_at,
                reverse=True,
            )
        ]


@app.post("/api/jobs")
async def create_job(
    file: UploadFile = File(...),
    language: str = Form("auto"),
    model_size: str = Form("large-v3"),
    translate_to: str = Form("none"),
    diarization: bool = Form(False),
    speaker_count: str = Form("2"),
    smart_split: bool = Form(True),
    ass_style: str = Form("standard"),
) -> Dict[str, object]:
    suffix = Path(file.filename or "").suffix.lower()
    if suffix not in ALLOWED_SUFFIXES:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported media format. Supported: {SUPPORTED_MEDIA_LABEL}.",
        )

    job_id = uuid.uuid4().hex
    safe_name = f"{job_id}{suffix}"
    audio_path = UPLOAD_DIR / safe_name

    with audio_path.open("wb") as buffer:
        shutil.copyfileobj(file.file, buffer)

    job = JobState(
        job_id=job_id,
        filename=safe_name,
        original_name=file.filename or safe_name,
        language=language,
        model_size=model_size,
        translate_to=translate_to,
        diarization=diarization,
        speaker_count=speaker_count,
        smart_split=smart_split,
        ass_style=ass_style,
    )

    with jobs_lock:
        jobs[job_id] = job
        snapshot = [asdict(item) for item in jobs.values()]
    persist_jobs(snapshot)

    executor.submit(
        transcribe_job,
        job_id,
        audio_path,
        language,
        model_size,
        translate_to,
        diarization,
        speaker_count,
        smart_split,
        ass_style,
    )
    return serialize_job(job)


@app.get("/api/jobs/{job_id}")
def get_job(job_id: str) -> Dict[str, object]:
    with jobs_lock:
        job = jobs.get(job_id)
        if not job:
            raise HTTPException(status_code=404, detail="Job not found.")
        return serialize_job(job)


@app.put("/api/jobs/{job_id}/draft")
def save_draft(job_id: str, payload: Dict[str, List[dict]]) -> Dict[str, object]:
    cleaned_segments = sanitize_segments(payload.get("segments", []))
    if not cleaned_segments:
        raise HTTPException(status_code=400, detail="Draft segments cannot be empty.")

    update_job(
        job_id,
        draft_segments=cleaned_segments,
        draft_updated_at=utc_now(),
        message="Draft auto-saved",
    )

    with jobs_lock:
        return serialize_job(jobs[job_id])


@app.put("/api/jobs/{job_id}/segments")
def update_segments(job_id: str, payload: Dict[str, List[dict]]) -> Dict[str, object]:
    incoming_segments = payload.get("segments", [])
    cleaned_segments = sanitize_segments(incoming_segments)
    if not cleaned_segments:
        raise HTTPException(status_code=400, detail="Segments cannot be empty.")

    with jobs_lock:
        job = jobs.get(job_id)
        if not job:
            raise HTTPException(status_code=404, detail="Job not found.")
        if not job.outputs:
            raise HTTPException(status_code=400, detail="Subtitle files are not ready yet.")
        source_language = job.detected_language or job.language
        translate_to = job.translate_to
        ass_style = job.ass_style
        output_paths = {name: Path(path) for name, path in job.outputs.items()}

    cleaned_segments = normalize_timeline(cleaned_segments)
    translated_segments = translate_segments(cleaned_segments, source_language, translate_to)
    translated_segments = write_all_outputs(translated_segments, output_paths, ass_style)
    update_job(
        job_id,
        segments=translated_segments,
        draft_segments=[],
        draft_updated_at=None,
        message="Subtitle timeline updated",
    )

    with jobs_lock:
        return serialize_job(jobs[job_id])


@app.delete("/api/jobs/{job_id}")
def delete_job(job_id: str) -> Dict[str, bool]:
    with jobs_lock:
        job = jobs.pop(job_id, None)
        if not job:
            raise HTTPException(status_code=404, detail="Job not found.")
        snapshot = [asdict(item) for item in jobs.values()]
    persist_jobs(snapshot)

    paths_to_remove = {UPLOAD_DIR / job.filename}
    paths_to_remove.update(Path(path) for path in job.outputs.values())

    for path in paths_to_remove:
        try:
            path.unlink(missing_ok=True)
        except OSError:
            continue

    return {"ok": True}


@app.get("/api/jobs/{job_id}/media")
def get_job_media(job_id: str) -> FileResponse:
    with jobs_lock:
        job = jobs.get(job_id)
        if not job:
            raise HTTPException(status_code=404, detail="Job not found.")
        media_path = UPLOAD_DIR / job.filename

    if not media_path.exists():
        raise HTTPException(status_code=404, detail="Source media not found.")

    media_type = mimetypes.guess_type(media_path.name)[0] or "application/octet-stream"
    return FileResponse(
        media_path,
        media_type=media_type,
        filename=job.original_name,
    )


@app.get("/api/jobs/{job_id}/download/{format_name}")
def download_output(job_id: str, format_name: str) -> FileResponse:
    with jobs_lock:
        job = jobs.get(job_id)
        if not job or not job.outputs:
            raise HTTPException(status_code=404, detail="Subtitle file not found.")
        if format_name not in job.outputs:
            raise HTTPException(status_code=404, detail="Requested format not found.")
        output_path = Path(job.outputs[format_name])

    media_types = {
        "srt": "application/x-subrip",
        "vtt": "text/vtt",
        "txt": "text/plain",
        "json": "application/json",
        "md": "text/markdown",
        "ass": "text/plain",
        "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    }

    return FileResponse(
        output_path,
        media_type=media_types.get(format_name, "application/octet-stream"),
        filename=output_path.name,
    )


@app.post("/api/jobs/{job_id}/save/{format_name}")
def save_output_to_downloads(job_id: str, format_name: str) -> Dict[str, str]:
    with jobs_lock:
        job = jobs.get(job_id)
        if not job or not job.outputs:
            raise HTTPException(status_code=404, detail="Subtitle file not found.")
        if format_name not in job.outputs:
            raise HTTPException(status_code=404, detail="Requested format not found.")
        output_path = Path(job.outputs[format_name])

    destination = copy_to_downloads(output_path)
    return {
        "ok": "true",
        "path": str(destination),
        "filename": destination.name,
        "directory": str(destination.parent),
    }


@app.get("/api/jobs/{job_id}/download-all")
def download_bundle(job_id: str) -> FileResponse:
    with jobs_lock:
        job = jobs.get(job_id)
        if not job or not job.outputs:
            raise HTTPException(status_code=404, detail="Subtitle bundle not found.")
        bundle_path = build_bundle_zip(job)

    return FileResponse(
        bundle_path,
        media_type="application/zip",
        filename=bundle_path.name,
    )


@app.post("/api/jobs/{job_id}/save-all")
def save_bundle_to_downloads(job_id: str) -> Dict[str, str]:
    with jobs_lock:
        job = jobs.get(job_id)
        if not job or not job.outputs:
            raise HTTPException(status_code=404, detail="Subtitle bundle not found.")
        bundle_path = build_bundle_zip(job)

    destination = copy_to_downloads(bundle_path)
    return {
        "ok": "true",
        "path": str(destination),
        "filename": destination.name,
        "directory": str(destination.parent),
    }
